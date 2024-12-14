from flask import Flask, render_template, request, jsonify, send_file, session
from openai import OpenAI
from dotenv import load_dotenv
import os
from werkzeug.utils import secure_filename
import tempfile
import PyPDF2
import io
import pandas as pd
from docx import Document
import xlrd

# Load environment variables
load_dotenv()

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx', 'csv', 'json', 'xlsx', 'xls'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.secret_key = os.urandom(24)  # Required for session management

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as e:
        print(f"Error reading PDF: {str(e)}")
        return None

def read_excel(file_path):
    try:
        # Try reading with pandas (works for xlsx and newer formats)
        df = pd.read_excel(file_path)
        return df.to_string()
    except Exception as e:
        try:
            # Fallback for older xls files
            workbook = xlrd.open_workbook(file_path)
            text = []
            for sheet in workbook.sheets():
                for row in range(sheet.nrows):
                    text.append('\t'.join(str(sheet.cell_value(row, col)) for col in range(sheet.ncols)))
            return '\n'.join(text)
        except Exception as e:
            print(f"Error reading Excel file: {str(e)}")
            return None

def read_word(file_path):
    try:
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                text.append('\t'.join(cell.text for cell in row.cells))
        return '\n'.join(text)
    except Exception as e:
        print(f"Error reading Word file: {str(e)}")
        return None

def read_file_content(file_path, file_type):
    if file_type == 'pdf':
        return read_pdf(file_path)
    elif file_type in ['xlsx', 'xls']:
        return read_excel(file_path)
    elif file_type in ['doc', 'docx']:
        return read_word(file_path)
    else:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                print(f"Error reading file: {str(e)}")
                return None

@app.route('/')
def home():
    try:
        return render_template('index.html')
    except Exception as e:
        print(f"Error loading template: {str(e)}")
        return f"Error: {str(e)}", 500

@app.route('/chat', methods=['POST'])
def chat():
    try:
        user_message = request.json['message']
        model = request.json.get('model', 'gpt-4o-mini')
        print(f"Using model for chat: {model}")
        
        # Initialize or reset model-specific conversation histories
        if 'conversation_histories' not in session:
            session['conversation_histories'] = {
                'gpt-4o-mini': [],
                'gpt-4o': []
            }
        
        # Get the conversation history for the current model
        current_history = session['conversation_histories'].get(model, [])
        
        # Add user message to model-specific history
        current_history.append({"role": "user", "content": user_message})
        
        # Create messages array with system prompt and history
        messages = [
            {"role": "system", "content": "Provide concise responses using markdown formatting. Keep responses brief but informative."}
        ] + current_history
        
        print(f"Sending messages to {model}:", messages)  # Debug print
        
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.5,
            max_tokens=2000
        )
        
        ai_response = response.choices[0].message.content
        if ai_response is None:
            return jsonify({"response": "Sorry, I couldn't generate a response"}), 500
        
        # Add AI response to model-specific history
        current_history.append({"role": "assistant", "content": ai_response})
        
        # Update the session with the new history
        session['conversation_histories'][model] = current_history
        session.modified = True  # Force session update
        
        print(f"Updated history for {model}:", session['conversation_histories'][model])  # Debug print
        
        return jsonify({"response": ai_response})
    
    except Exception as e:
        print("Error:", str(e))
        return jsonify({"response": f"Error: {str(e)}"}), 500

@app.route('/clear-history', methods=['POST'])
def clear_history():
    try:
        model = request.json.get('model')
        print(f"Clearing history for model: {model}")  # Debug print
        
        if 'conversation_histories' in session:
            if model:
                # Clear history for specific model
                session['conversation_histories'][model] = []
            else:
                # Clear all histories
                session['conversation_histories'] = {
                    'gpt-4o-mini': [],
                    'gpt-4o': []
                }
            session.modified = True  # Force session update
            
        print("Current session histories:", session.get('conversation_histories'))  # Debug print
        return jsonify({"status": "success"})
    except Exception as e:
        print("Error clearing history:", str(e))
        return jsonify({"error": str(e)}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400
        
        file = request.files['file']
        message = request.form.get('message', '')
        model = request.form.get('model', 'gpt-4o-mini')
        print(f"Using model for file: {model}")
        
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Get file extension
            file_type = filename.rsplit('.', 1)[1].lower()
            
            # Read file content based on type
            file_content = read_file_content(filepath, file_type)
            
            if file_content is None:
                os.remove(filepath)
                return jsonify({"error": f"Could not read {file_type.upper()} file content"}), 400
            
            # Initialize model-specific conversation histories if they don't exist
            if 'conversation_histories' not in session:
                session['conversation_histories'] = {
                    'gpt-4o-mini': [],
                    'gpt-4o': []
                }
            
            # Get the conversation history for the current model
            current_history = session['conversation_histories'].get(model, [])
            
            # Prepare the prompt with file type context
            prompt = f"This is content from a {file_type.upper()} file:\n\n{file_content}\n\n"
            if message:
                prompt += f"User's request: {message}\n\n"
            prompt += "Please analyze this content and provide insights."
            
            # Add file content and message to model-specific history
            current_history.append({"role": "user", "content": prompt})
            
            try:
                # Send to OpenAI with conversation history
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "Provide concise responses using markdown formatting. Keep responses brief but informative."}
                    ] + current_history,
                    temperature=0.5,
                    max_tokens=4000
                )
                
                ai_response = response.choices[0].message.content
                
                # Add AI response to model-specific history
                current_history.append({"role": "assistant", "content": ai_response})
                
                # Update the session with the new history
                session['conversation_histories'][model] = current_history
                
                # Limit history length
                if len(current_history) > 10:
                    session['conversation_histories'][model] = current_history[-10:]
                
                os.remove(filepath)  # Clean up the file
                return jsonify({"response": ai_response})
            except Exception as e:
                os.remove(filepath)
                return jsonify({"error": f"AI processing error: {str(e)}"}), 500
        
        return jsonify({"error": "File type not allowed"}), 400
    except Exception as e:
        print(f"Upload error: {str(e)}")
        if 'filepath' in locals() and os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({"error": f"Error processing file: {str(e)}"}), 500

@app.route('/download', methods=['POST'])
def download_file():
    try:
        content = request.json.get('content')
        filename = request.json.get('filename', 'response.txt')
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(mode='w', delete=False) as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name
        
        return send_file(
            temp_path,
            as_attachment=True,
            download_name=filename,
            mimetype='text/plain'
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=True)
